# -*- coding: utf-8 -*-
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sklearn.ensemble import RandomForestRegressor
from sklearn.linear_model import LinearRegression
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
from sklearn.model_selection import train_test_split
from sklearn.neighbors import KNeighborsRegressor
from sklearn.pipeline import make_pipeline
from sklearn.preprocessing import StandardScaler
from sklearn.tree import DecisionTreeRegressor


BASE = Path(__file__).resolve().parent
ASSETS = BASE / "report_assets"
OUTDIR = Path(r"C:/Users/human-21/Documents/Codex/2026-06-18/new-chat/outputs")
REPORT = OUTDIR / "car_mpg_model_comparison_report.docx"

FEATURES = [
    "cylinders",
    "displacement",
    "horsepower",
    "weight",
    "acceleration",
    "model_year",
]

LABELS = {
    "cylinders": "실린더 수",
    "displacement": "배기량",
    "horsepower": "마력",
    "weight": "무게",
    "acceleration": "가속력",
    "model_year": "연식",
}


def build_models():
    return {
        "선형회귀": LinearRegression(),
        "의사결정나무": DecisionTreeRegressor(random_state=42, max_depth=5),
        "랜덤포레스트": RandomForestRegressor(
            random_state=42,
            n_estimators=300,
            max_depth=8,
        ),
        "KNN 회귀": make_pipeline(
            StandardScaler(),
            KNeighborsRegressor(n_neighbors=5),
        ),
    }


def style_run(run, size=10.5, bold=False):
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    style_run(run, 9.5, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell_text(cell, value, row_index == 0 or col_index == 0)
            if row_index == 0 or col_index == 0:
                shade(cell, "D9EAF7")
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        style_run(run, 16 if level == 1 else 13, True)
        run.font.color.rgb = RGBColor(22, 32, 51)
    return paragraph


def add_paragraph(doc, text=""):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    style_run(run, 10.5)
    paragraph.paragraph_format.line_spacing = 1.35
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    style_run(run, 10)
    return paragraph


def train_and_compare(df):
    X = df[FEATURES]
    y = df["mpg"]
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.2, random_state=42
    )

    trained_models = {}
    predictions = {}
    rows = []

    for name, model in build_models().items():
        model.fit(X_train, y_train)
        y_pred = model.predict(X_test)
        rmse = np.sqrt(mean_squared_error(y_test, y_pred))
        rows.append(
            {
                "모델": name,
                "R2": r2_score(y_test, y_pred),
                "MAE": mean_absolute_error(y_test, y_pred),
                "RMSE": rmse,
            }
        )
        trained_models[name] = model
        predictions[name] = y_pred

    results = pd.DataFrame(rows).sort_values("RMSE").reset_index(drop=True)
    best_name = results.loc[0, "모델"]
    return trained_models, predictions, results, best_name, X_test, y_test


def get_importance(model):
    estimator = model
    if hasattr(model, "named_steps"):
        estimator = list(model.named_steps.values())[-1]

    if hasattr(estimator, "feature_importances_"):
        values = estimator.feature_importances_
    elif hasattr(estimator, "coef_"):
        values = np.abs(estimator.coef_)
    else:
        values = np.zeros(len(FEATURES))

    importance = pd.DataFrame(
        {
            "특성": [LABELS[item] for item in FEATURES],
            "중요도": values,
        }
    ).sort_values("중요도", ascending=False)
    return importance


def create_charts(df, predictions, results, best_name, y_test, importance):
    ASSETS.mkdir(exist_ok=True)
    plt.rcParams["font.family"] = ["Malgun Gothic"]
    plt.rcParams["axes.unicode_minus"] = False

    fig, ax = plt.subplots(figsize=(7, 4))
    ax.hist(df["mpg"], bins=18, color="#1f7a8c", edgecolor="white")
    ax.set_title("자동차 연비(MPG) 분포")
    ax.set_xlabel("MPG")
    ax.set_ylabel("빈도")
    fig.tight_layout()
    mpg_hist = ASSETS / "model_report_mpg_distribution.png"
    fig.savefig(mpg_hist, dpi=160)
    plt.close(fig)

    corr = df[["mpg", *FEATURES]].corr()
    fig, ax = plt.subplots(figsize=(7, 5))
    image = ax.imshow(corr, cmap="RdBu_r", vmin=-1, vmax=1)
    ax.set_xticks(range(len(corr.columns)))
    ax.set_yticks(range(len(corr.columns)))
    ax.set_xticklabels(["mpg", *[LABELS[item] for item in FEATURES]], rotation=45, ha="right")
    ax.set_yticklabels(["mpg", *[LABELS[item] for item in FEATURES]])
    for row in range(len(corr.columns)):
        for col in range(len(corr.columns)):
            ax.text(col, row, f"{corr.iloc[row, col]:.2f}", ha="center", va="center", fontsize=8)
    fig.colorbar(image, ax=ax, shrink=0.8)
    ax.set_title("변수 간 상관관계 Heatmap")
    fig.tight_layout()
    heatmap = ASSETS / "model_report_correlation_heatmap.png"
    fig.savefig(heatmap, dpi=160)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(7, 4))
    ordered = results.sort_values("RMSE", ascending=True)
    colors = ["#1f7a8c" if name == best_name else "#8aa4b8" for name in ordered["모델"]]
    ax.barh(ordered["모델"], ordered["RMSE"], color=colors)
    ax.invert_yaxis()
    ax.set_title("모델별 RMSE 비교")
    ax.set_xlabel("RMSE")
    fig.tight_layout()
    rmse_chart = ASSETS / "model_report_rmse_comparison.png"
    fig.savefig(rmse_chart, dpi=160)
    plt.close(fig)

    fig, axes = plt.subplots(2, 3, figsize=(10, 6.5))
    axes = axes.flatten()
    min_value = min(y_test.min(), min(pred.min() for pred in predictions.values()))
    max_value = max(y_test.max(), max(pred.max() for pred in predictions.values()))
    for index, (name, y_pred) in enumerate(predictions.items()):
        ax = axes[index]
        ax.scatter(y_test, y_pred, alpha=0.7, color="#1769aa", s=22)
        ax.plot([min_value, max_value], [min_value, max_value], "--", color="#b3261e")
        ax.set_title(name)
        ax.set_xlabel("실제 MPG")
        ax.set_ylabel("예측 MPG")
    axes[-1].axis("off")
    fig.suptitle("모델별 실제 연비와 예측 연비 비교", fontsize=14)
    fig.tight_layout()
    prediction_chart = ASSETS / "model_report_prediction_comparison.png"
    fig.savefig(prediction_chart, dpi=160)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(7, 4))
    ax.barh(importance["특성"], importance["중요도"], color="#1f7a8c")
    ax.invert_yaxis()
    ax.set_title(f"{best_name} 특성 영향도")
    ax.set_xlabel("중요도")
    fig.tight_layout()
    importance_chart = ASSETS / "model_report_best_feature_importance.png"
    fig.savefig(importance_chart, dpi=160)
    plt.close(fig)

    return mpg_hist, heatmap, rmse_chart, prediction_chart, importance_chart


def main():
    OUTDIR.mkdir(parents=True, exist_ok=True)
    df = pd.read_csv(BASE / "data" / "mpg.csv")
    df = df[["mpg", *FEATURES]].dropna()

    trained_models, predictions, results, best_name, X_test, y_test = train_and_compare(df)
    importance = get_importance(trained_models[best_name])
    mpg_hist, heatmap, rmse_chart, prediction_chart, importance_chart = create_charts(
        df, predictions, results, best_name, y_test, importance
    )

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("2026년도 ML 모델 개발 과제\nAI개발 수행내역서")
    style_run(run, 22, True)
    doc.add_paragraph("")
    add_table(
        doc,
        [
            ["과제명", "자동차 제원 데이터를 활용한 연비 예측 모델 비교 및 웹 시각화"],
            ["담당자", ""],
            ["작성일", "2026년 6월 22일"],
        ],
    )
    doc.add_paragraph("")
    add_paragraph(
        doc,
        "본 문서는 자동차 제원 데이터를 활용하여 여러 회귀 모델을 학습시킨 뒤 성능을 비교하고, "
        "가장 예측 성능이 좋은 모델을 Flask 웹앱에 적용한 과정을 정리한 AI개발 수행내역서이다.",
    )
    doc.add_page_break()

    add_heading(doc, "AI개발 수행내용", 1)
    add_heading(doc, "1. 사업과제", 2)
    add_paragraph(doc, "자동차 제원 데이터를 활용한 AI 기반 연비 예측 모델 비교 및 웹 프로토타입 개발")

    add_heading(doc, "2. 개요 및 현황", 2)
    add_heading(doc, "2.1 추진배경 및 목적", 3)
    for item in [
        "자동차 연비는 차량 유지비, 에너지 소비량, 환경 부담과 직접적으로 연결되는 중요한 지표이다.",
        "같은 데이터라도 모델 종류에 따라 예측 성능이 달라질 수 있으므로 여러 회귀 모델을 비교하는 과정이 필요하다.",
        "본 프로젝트는 선형회귀, 의사결정나무, 랜덤포레스트, KNN 회귀 모델을 학습시켜 성능을 비교하였다.",
        "비교 결과 RMSE가 가장 낮은 모델을 최종 예측 모델로 선택하고, Flask 웹 화면에서 사용자 입력값에 대한 예상 연비를 출력하도록 구현하였다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.2 과제 범위", 3)
    add_table(
        doc,
        [
            ["과제구분", "내용"],
            ["데이터", "Auto MPG 공개 데이터셋 수집 및 전처리"],
            ["모델 학습", "선형회귀, 의사결정나무, 랜덤포레스트, KNN 회귀 학습"],
            ["모델 평가", "R², MAE, RMSE를 기준으로 모델별 성능 비교"],
            ["최종 모델", f"RMSE 기준 최종 선택 모델: {best_name}"],
            ["웹 프로토타입", "Flask 입력 화면, 최종 모델 예측, 모델별 성능 비교 표시"],
            ["DB 연동", "MariaDB에 예측 입력값과 예측 결과 저장"],
        ],
    )

    add_heading(doc, "2.3 과제 추진 방법", 3)
    add_table(
        doc,
        [
            ["단계", "수행 내용"],
            ["DATA IMPORTING", "mpg.csv 데이터 로드 및 주요 변수 확인"],
            ["MODEL TRAINING", "동일한 학습/테스트 데이터로 여러 회귀 모델 학습"],
            ["MODEL COMPARISON", "R², MAE, RMSE 기준으로 모델별 성능 비교"],
            ["AUTO PREDICTION", "최고 성능 모델을 Flask 웹앱 예측 모델로 사용"],
        ],
    )
    doc.add_page_break()

    add_heading(doc, "연구개발 주요 결과물", 1)
    add_heading(doc, "1. 데이터 수집", 2)
    add_paragraph(
        doc,
        "본 프로젝트에서는 Seaborn 예제 데이터로 제공되는 Auto MPG 데이터셋을 사용하였다. "
        "종속변수는 자동차 연비를 의미하는 mpg이며, 독립변수는 실린더 수, 배기량, 마력, 무게, 가속력, 연식이다.",
    )
    add_table(
        doc,
        [
            ["항목", "내용"],
            ["데이터명", "Auto MPG Dataset"],
            ["출처", "UCI Machine Learning Repository / Seaborn mpg.csv"],
            ["데이터 개수", f"{len(df)}개"],
            ["종속변수(Y)", "mpg: 자동차 연비"],
            ["독립변수(X)", "cylinders, displacement, horsepower, weight, acceleration, model_year"],
        ],
    )

    add_heading(doc, "2. 데이터 분석 및 전처리", 2)
    summary_rows = [["변수", "평균", "최소값", "최대값"]]
    for column in ["mpg", *FEATURES]:
        summary_rows.append(
            [
                LABELS.get(column, column),
                f"{df[column].mean():.2f}",
                f"{df[column].min():.2f}",
                f"{df[column].max():.2f}",
            ]
        )
    add_table(doc, summary_rows)
    add_paragraph(
        doc,
        "결측치가 있는 행은 제거하였고, 모든 모델이 동일한 데이터 조건에서 비교될 수 있도록 동일한 train_test_split 설정을 사용하였다.",
    )
    doc.add_picture(str(mpg_hist), width=Inches(5.8))
    add_paragraph(doc, "연비 데이터는 약 10~45 MPG 범위에 분포하며, 중간 연비 구간에 데이터가 집중되어 있다.")
    doc.add_picture(str(heatmap), width=Inches(5.8))
    add_paragraph(
        doc,
        "상관관계 분석 결과 차량 무게, 배기량, 마력, 실린더 수는 연비와 음의 상관관계를 보였다.",
    )
    doc.add_page_break()

    add_heading(doc, "3. 모델 학습 및 모델별 성능 비교", 2)
    add_heading(doc, "3.1 학습 모델 정의", 3)
    add_table(
        doc,
        [
            ["모델", "특징"],
            ["선형회귀", "독립변수와 종속변수 사이의 선형 관계를 학습하는 기본 회귀 모델"],
            ["의사결정나무", "조건 분기를 통해 데이터를 나누며 예측하는 트리 기반 모델"],
            ["랜덤포레스트", "여러 개의 결정나무를 앙상블하여 예측 안정성을 높인 모델"],
            ["KNN 회귀", "입력값과 가까운 이웃 데이터들의 평균을 활용해 예측하는 모델"],
        ],
    )

    add_heading(doc, "3.2 모델별 성능 평가", 3)
    rows = [["순위", "모델", "R²", "MAE", "RMSE"]]
    for index, row in results.iterrows():
        rows.append(
            [
                index + 1,
                row["모델"],
                f"{row['R2']:.3f}",
                f"{row['MAE']:.3f}",
                f"{row['RMSE']:.3f}",
            ]
        )
    add_table(doc, rows)
    doc.add_picture(str(rmse_chart), width=Inches(5.8))
    add_paragraph(
        doc,
        f"모델 비교 결과 RMSE가 가장 낮은 모델은 {best_name}으로 나타났다. "
        "따라서 본 프로젝트의 Flask 웹앱에서는 해당 모델을 최종 예측 모델로 사용하였다.",
    )

    add_heading(doc, "3.3 모델별 예측값 그래프", 3)
    doc.add_picture(str(prediction_chart), width=Inches(6.3))
    add_paragraph(
        doc,
        "각 그래프는 테스트 데이터의 실제 MPG와 모델이 예측한 MPG를 비교한 것이다. "
        "점들이 대각선에 가까울수록 실제값과 예측값이 잘 맞는 것으로 해석할 수 있다.",
    )

    add_heading(doc, "3.4 최종 모델 특성 영향 분석", 3)
    rows = [["특성", "중요도"]]
    for _, row in importance.iterrows():
        rows.append([row["특성"], f"{row['중요도']:.4f}"])
    add_table(doc, rows)
    doc.add_picture(str(importance_chart), width=Inches(5.8))
    add_paragraph(
        doc,
        f"최종 선택 모델인 {best_name}의 특성 영향도를 확인한 결과, 모델은 차량 제원 중 일부 변수를 더 크게 반영하여 연비를 예측하였다.",
    )
    doc.add_page_break()

    add_heading(doc, "4. 프로토타입 구현", 2)
    add_heading(doc, "4.1 Flask 웹 화면", 3)
    add_paragraph(
        doc,
        "Flask 웹앱은 사용자가 차량 제원을 입력하면 최고 성능 모델을 사용하여 예상 연비를 출력한다. "
        "화면에는 최종 선택 모델, 모델별 성능 비교표, MariaDB 연결 상태, 최근 예측 기록도 함께 표시된다.",
    )
    screenshot = ASSETS / "flask_screen.png"
    if screenshot.exists():
        doc.add_picture(str(screenshot), width=Inches(6.0))

    add_heading(doc, "4.2 MariaDB 연동", 3)
    add_paragraph(
        doc,
        "MariaDB는 예측 요청 기록을 저장하기 위해 사용하였다. 사용자가 예측 버튼을 누르면 입력값과 예측 MPG가 predictions 테이블에 저장된다.",
    )
    add_table(
        doc,
        [
            ["구분", "내용"],
            ["DBMS", "MariaDB"],
            ["Database", "project1_car"],
            ["Table", "predictions"],
            ["저장 데이터", "입력 차량 제원, 예측 MPG, 생성 시간"],
            ["연결 방식", "db.py에서 PyMySQL을 이용해 연결"],
        ],
    )

    add_heading(doc, "5. 결론", 2)
    for item in [
        "본 프로젝트에서는 자동차 제원 데이터를 활용하여 여러 회귀 모델의 연비 예측 성능을 비교하였다.",
        f"성능 비교 결과 RMSE 기준으로 {best_name} 모델이 가장 좋은 예측 성능을 보였다.",
        "최종 모델을 Flask 웹앱에 적용하여 사용자가 직접 자동차 정보를 입력하고 예상 연비를 확인할 수 있도록 구현하였다.",
        "MariaDB 연동을 통해 예측 입력값과 결과를 저장하여 단순 모델 실습을 웹 기반 데이터 관리 시스템으로 확장하였다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. 향후 발전 방향", 2)
    for item in [
        "제조사, 차량명, 국가(origin) 등 범주형 변수를 추가하면 더 정교한 예측이 가능하다.",
        "모델별 하이퍼파라미터 튜닝을 추가하면 랜덤포레스트나 KNN 회귀의 성능을 더 개선할 수 있다.",
        "MariaDB에 저장된 예측 기록을 활용해 사용자 입력 패턴과 예측 결과 변화도 시각화할 수 있다.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "참고자료", 1)
    for item in [
        "UCI Machine Learning Repository, Auto MPG Dataset: https://archive.ics.uci.edu/dataset/9/auto+mpg",
        "Seaborn Data Repository, mpg.csv: https://raw.githubusercontent.com/mwaskom/seaborn-data/master/mpg.csv",
        "Flask Documentation: https://flask.palletsprojects.com/",
        "scikit-learn Documentation: https://scikit-learn.org/",
        "MariaDB Documentation: https://mariadb.org/documentation/",
    ]:
        add_bullet(doc, item)

    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    main()
