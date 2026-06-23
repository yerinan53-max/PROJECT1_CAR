# -*- coding: utf-8 -*-
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sklearn.ensemble import RandomForestRegressor
from sklearn.linear_model import LinearRegression
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
from sklearn.model_selection import train_test_split
from sklearn.neighbors import KNeighborsRegressor
from sklearn.pipeline import make_pipeline
from sklearn.preprocessing import StandardScaler
from sklearn.tree import DecisionTreeRegressor


BASE_DIR = Path(__file__).resolve().parent
DATA_PATH = BASE_DIR / "data" / "mpg.csv"
ASSET_DIR = BASE_DIR / "report_assets"
REPORT_PATH = BASE_DIR / "자동차_연비_예측_최종보고서.docx"
FALLBACK_REPORT_PATH = BASE_DIR / "자동차_연비_예측_최종보고서_그래프포함.docx"

FEATURES = [
    "cylinders",
    "displacement",
    "horsepower",
    "weight",
    "acceleration",
    "model_year",
]

FEATURE_LABELS = {
    "mpg": "연비(MPG)",
    "cylinders": "실린더 수",
    "displacement": "배기량",
    "horsepower": "마력",
    "weight": "무게",
    "acceleration": "가속력",
    "model_year": "연식",
}


def build_models():
    return {
        "선형회귀": LinearRegression(),
        "의사결정나무": DecisionTreeRegressor(random_state=42, max_depth=5),
        "랜덤 포레스트": RandomForestRegressor(
            random_state=42,
            n_estimators=300,
            max_depth=8,
        ),
        "KNN 회귀": make_pipeline(
            StandardScaler(),
            KNeighborsRegressor(n_neighbors=5),
        ),
    }


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, fill=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    set_run_font(run, size=9.3, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(text)
    size = 16 if level == 1 else 13 if level == 2 else 11.5
    set_run_font(run, size=size, bold=True, color=(22, 32, 51))
    return paragraph


def add_paragraph(doc, text=""):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run)
    paragraph.paragraph_format.line_spacing = 1.35
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run, size=10)
    return paragraph


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, fill="D9EAF7")

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)

    doc.add_paragraph("")
    return table


def add_image_with_caption(doc, image_path, caption, width=5.9):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))

    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=9.2, color=(77, 95, 120))


def load_data():
    df = pd.read_csv(DATA_PATH)
    return df[["mpg", *FEATURES]].dropna()


def train_models(df):
    x = df[FEATURES]
    y = df["mpg"]
    x_train, x_test, y_train, y_test = train_test_split(
        x, y, test_size=0.2, random_state=42
    )

    trained_models = {}
    predictions = {}
    metric_rows = []

    for name, model in build_models().items():
        model.fit(x_train, y_train)
        y_pred = model.predict(x_test)
        rmse = float(np.sqrt(mean_squared_error(y_test, y_pred)))
        metric_rows.append(
            {
                "모델": name,
                "R2": float(r2_score(y_test, y_pred)),
                "MAE": float(mean_absolute_error(y_test, y_pred)),
                "RMSE": rmse,
            }
        )
        trained_models[name] = model
        predictions[name] = y_pred

    results = pd.DataFrame(metric_rows).sort_values("RMSE").reset_index(drop=True)
    best_model_name = results.loc[0, "모델"]
    return trained_models, predictions, results, best_model_name, x_train, x_test, y_test


def get_feature_importance(model):
    estimator = model
    if hasattr(model, "named_steps"):
        estimator = list(model.named_steps.values())[-1]

    if hasattr(estimator, "feature_importances_"):
        values = estimator.feature_importances_
    elif hasattr(estimator, "coef_"):
        values = np.abs(estimator.coef_)
    else:
        values = np.zeros(len(FEATURES))

    return (
        pd.DataFrame(
            {
                "특성": [FEATURE_LABELS[feature] for feature in FEATURES],
                "중요도": values,
            }
        )
        .sort_values("중요도", ascending=False)
        .reset_index(drop=True)
    )


def setup_plot_style():
    plt.rcParams["font.family"] = ["Malgun Gothic", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.facecolor"] = "white"
    plt.rcParams["axes.facecolor"] = "white"


def save_mpg_distribution(df):
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    ax.hist(df["mpg"], bins=18, color="#1f7a8c", edgecolor="white")
    ax.set_title("자동차 연비(MPG) 분포")
    ax.set_xlabel("MPG")
    ax.set_ylabel("데이터 수")
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    path = ASSET_DIR / "final_mpg_distribution.png"
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def save_correlation_heatmap(df):
    corr = df[["mpg", *FEATURES]].corr()
    labels = [FEATURE_LABELS[column] for column in corr.columns]

    fig, ax = plt.subplots(figsize=(7.2, 5.5))
    image = ax.imshow(corr, cmap="RdBu_r", vmin=-1, vmax=1)
    ax.set_xticks(range(len(labels)))
    ax.set_yticks(range(len(labels)))
    ax.set_xticklabels(labels, rotation=35, ha="right")
    ax.set_yticklabels(labels)

    for row in range(len(labels)):
        for column in range(len(labels)):
            ax.text(
                column,
                row,
                f"{corr.iloc[row, column]:.2f}",
                ha="center",
                va="center",
                fontsize=8,
                color="#162033",
            )

    fig.colorbar(image, ax=ax, shrink=0.78)
    ax.set_title("변수 간 상관관계")
    fig.tight_layout()
    path = ASSET_DIR / "final_correlation_heatmap.png"
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def save_rmse_comparison(results, best_model_name):
    ordered = results.sort_values("RMSE", ascending=True)
    colors = ["#1f7a8c" if name == best_model_name else "#9fb4c7" for name in ordered["모델"]]

    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    ax.barh(ordered["모델"], ordered["RMSE"], color=colors)
    ax.invert_yaxis()
    ax.set_title("모델별 RMSE 비교")
    ax.set_xlabel("RMSE")
    ax.grid(axis="x", alpha=0.25)

    for index, value in enumerate(ordered["RMSE"]):
        ax.text(value + 0.03, index, f"{value:.3f}", va="center", fontsize=9)

    fig.tight_layout()
    path = ASSET_DIR / "final_rmse_comparison.png"
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def save_prediction_comparison(predictions, y_test):
    min_value = min(y_test.min(), *(prediction.min() for prediction in predictions.values()))
    max_value = max(y_test.max(), *(prediction.max() for prediction in predictions.values()))

    fig, axes = plt.subplots(2, 2, figsize=(9.2, 7.0))
    axes = axes.flatten()
    for ax, (name, y_pred) in zip(axes, predictions.items()):
        ax.scatter(y_test, y_pred, color="#1769aa", alpha=0.72, s=24)
        ax.plot([min_value, max_value], [min_value, max_value], "--", color="#b3261e")
        ax.set_title(name)
        ax.set_xlabel("실제 MPG")
        ax.set_ylabel("예측 MPG")
        ax.grid(alpha=0.22)

    fig.suptitle("모델별 실제값과 예측값 비교", fontsize=14)
    fig.tight_layout()
    path = ASSET_DIR / "final_prediction_comparison.png"
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def save_feature_importance(importance, best_model_name):
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    ax.barh(importance["특성"], importance["중요도"], color="#1f7a8c")
    ax.invert_yaxis()
    ax.set_title(f"{best_model_name} 특성 영향도")
    ax.set_xlabel("중요도")
    ax.grid(axis="x", alpha=0.25)

    for index, value in enumerate(importance["중요도"]):
        ax.text(value + 0.003, index, f"{value:.4f}", va="center", fontsize=9)

    fig.tight_layout()
    path = ASSET_DIR / "final_feature_importance.png"
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def create_charts(df, predictions, results, best_model_name, y_test, importance):
    ASSET_DIR.mkdir(exist_ok=True)
    setup_plot_style()
    return {
        "mpg": save_mpg_distribution(df),
        "correlation": save_correlation_heatmap(df),
        "rmse": save_rmse_comparison(results, best_model_name),
        "prediction": save_prediction_comparison(predictions, y_test),
        "importance": save_feature_importance(importance, best_model_name),
    }


def add_title_page(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("자동차 연비 예측 프로젝트 최종 보고서")
    set_run_font(run, size=20, bold=True, color=(21, 76, 121))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Auto MPG 데이터 기반 회귀 모델 비교 및 Streamlit 앱 구현")
    set_run_font(run, size=11, color=(77, 95, 120))
    doc.add_paragraph("")


def create_report(df, results, best_model_name, x_train, x_test, charts, importance):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "맑은 고딕"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    styles["Normal"].font.size = Pt(10.5)

    add_title_page(doc)

    add_heading(doc, "1. 프로젝트 개요", 1)
    add_paragraph(
        doc,
        "본 프로젝트는 자동차의 실린더 수, 배기량, 마력, 무게, 가속력, 연식을 이용하여 "
        "예상 연비(MPG)를 예측하는 머신러닝 프로젝트이다. 여러 회귀 모델을 동일한 데이터로 학습시킨 뒤 "
        "성능을 비교하고, 가장 좋은 모델을 Streamlit 웹 앱에 적용하였다.",
    )

    add_heading(doc, "2. 제출 산출물 구성", 1)
    add_table(
        doc,
        ["구분", "파일", "설명"],
        [
            ["모델 학습 파일", "car_mpg_training.ipynb", "데이터 로드, 전처리, 모델 학습, 성능 비교 과정을 정리한 Jupyter Notebook"],
            ["Streamlit 앱", "app.py", "사용자 입력값으로 예상 MPG를 계산하는 웹 앱"],
            ["의존성 파일", "requirements.txt", "Streamlit Cloud 배포에 필요한 라이브러리 목록"],
            ["최종 보고서", "자동차_연비_예측_최종보고서.docx", "학습 과정, 그래프, 결과 해석을 정리한 한글 보고서"],
        ],
    )

    add_heading(doc, "3. 데이터 설명", 1)
    add_paragraph(
        doc,
        "데이터는 Auto MPG 공개 데이터셋을 사용하였다. 결측치가 있는 행은 제거하고, "
        "연비 예측에 직접적으로 사용한 수치형 변수만 선택하였다.",
    )
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["데이터셋", "Auto MPG Dataset"],
            ["데이터 개수", f"{len(df)}개"],
            ["예측 대상", "mpg: 자동차 연비"],
            ["입력 변수", ", ".join(FEATURES)],
            ["학습/테스트 분리", f"학습 {len(x_train)}개, 테스트 {len(x_test)}개"],
        ],
    )

    summary_rows = []
    for column in ["mpg", *FEATURES]:
        summary_rows.append(
            [
                FEATURE_LABELS[column],
                f"{df[column].mean():.2f}",
                f"{df[column].min():.2f}",
                f"{df[column].max():.2f}",
            ]
        )
    add_table(doc, ["변수", "평균", "최소", "최대"], summary_rows)

    add_image_with_caption(
        doc,
        charts["mpg"],
        "그림 1. 자동차 연비(MPG) 분포",
    )
    add_paragraph(
        doc,
        "연비 분포 그래프를 보면 대부분의 차량이 중간 연비 구간에 몰려 있다. "
        "데이터의 범위와 집중 구간을 확인하면 모델이 어떤 값의 예측을 많이 학습하게 되는지 이해할 수 있다.",
    )

    add_image_with_caption(
        doc,
        charts["correlation"],
        "그림 2. 입력 변수와 연비 사이의 상관관계",
    )
    add_paragraph(
        doc,
        "상관관계 그래프에서는 무게, 배기량, 마력, 실린더 수가 연비와 음의 상관관계를 보인다. "
        "즉 차량이 무겁거나 엔진 규모가 클수록 MPG가 낮아지는 경향이 나타난다.",
    )

    doc.add_page_break()
    add_heading(doc, "4. 모델 학습 과정", 1)
    add_paragraph(
        doc,
        "모델 학습은 car_mpg_training.ipynb에서 수행하였다. 수업 범위에 맞추어 선형회귀, "
        "의사결정나무, 랜덤 포레스트, KNN 회귀 4개 모델만 비교하였다.",
    )
    add_table(
        doc,
        ["모델", "설명"],
        [
            ["선형회귀", "입력 변수와 연비 사이의 선형 관계를 학습하는 기본 회귀 모델"],
            ["의사결정나무", "조건 분기를 통해 데이터를 나누고 예측하는 트리 기반 모델"],
            ["랜덤 포레스트", "여러 개의 결정나무를 함께 사용하여 예측 안정성을 높인 모델"],
            ["KNN 회귀", "입력값과 가까운 이웃 데이터들의 평균을 이용해 예측하는 모델"],
        ],
    )

    add_heading(doc, "5. 학습 모델별 성능 비교", 1)
    add_paragraph(
        doc,
        "모델 성능은 R², MAE, RMSE를 기준으로 비교하였다. RMSE는 예측 오차를 의미하므로 낮을수록 좋고, "
        "R²는 모델 설명력을 의미하므로 높을수록 좋다.",
    )

    result_rows = []
    for index, row in results.iterrows():
        result_rows.append(
            [
                index + 1,
                row["모델"],
                f"{row['R2']:.3f}",
                f"{row['MAE']:.3f}",
                f"{row['RMSE']:.3f}",
                "최종 모델" if row["모델"] == best_model_name else "비교 모델",
            ]
        )
    add_table(doc, ["순위", "모델", "R²", "MAE", "RMSE", "선택 여부"], result_rows)

    add_image_with_caption(
        doc,
        charts["rmse"],
        "그림 3. 모델별 RMSE 비교",
    )
    add_paragraph(
        doc,
        f"RMSE 비교 결과 {best_model_name} 모델이 가장 낮은 오차를 보였다. "
        "따라서 Streamlit 앱에서는 이 모델을 최종 예측 모델로 사용하였다.",
    )

    add_image_with_caption(
        doc,
        charts["prediction"],
        "그림 4. 학습 모델별 실제값과 예측값 비교",
        width=6.4,
    )
    add_paragraph(
        doc,
        "실제값과 예측값 비교 그래프에서 점들이 붉은 대각선에 가까울수록 예측이 정확하다고 볼 수 있다. "
        "모델별 그래프를 함께 보면 단순 성능표뿐 아니라 예측값이 실제값을 얼마나 잘 따라가는지도 확인할 수 있다.",
    )

    add_heading(doc, "6. 최종 모델 특성 영향 분석", 1)
    add_paragraph(
        doc,
        "최종 선택 모델 기준으로 각 입력 변수가 예측에 어느 정도 반영되었는지 확인하였다. "
        "중요도가 높은 변수일수록 연비 예측 결과에 더 큰 영향을 준다.",
    )
    add_table(
        doc,
        ["특성", "중요도"],
        [[row["특성"], f"{row['중요도']:.4f}"] for _, row in importance.iterrows()],
    )
    add_image_with_caption(
        doc,
        charts["importance"],
        "그림 5. 최종 모델의 특성 영향도",
    )
    add_paragraph(
        doc,
        "특성 영향도 그래프를 통해 모델이 어떤 차량 제원을 더 중요하게 판단하는지 확인할 수 있다. "
        "보고서에서는 숫자 표와 그래프를 함께 제시하여 결과를 더 쉽게 비교할 수 있도록 하였다.",
    )

    doc.add_page_break()
    add_heading(doc, "7. Streamlit 앱 구현", 1)
    add_paragraph(
        doc,
        "웹 앱은 Streamlit으로 구현하였다. 사용자가 차량 정보를 입력하고 버튼을 누르면 최종 모델이 예상 MPG를 계산한다. "
        "화면에는 모델 성능, 모델별 비교, 최근 예측 기록, 특성 영향 분석이 함께 표시된다.",
    )
    add_bullet(doc, "입력값: 실린더 수, 배기량, 마력, 무게, 가속력, 연식")
    add_bullet(doc, "출력값: 예상 연비(MPG), 최종 선택 모델, 성능 지표")
    add_bullet(doc, "Streamlit Cloud 배포를 고려하여 MariaDB 저장 기능은 기본 비활성화하였다.")
    add_bullet(doc, "로컬 환경에서 ENABLE_DB=true를 설정하면 MariaDB 저장 기능을 사용할 수 있다.")

    add_heading(doc, "8. 결론 및 향후 개선 방향", 1)
    add_paragraph(
        doc,
        "본 프로젝트는 Auto MPG 데이터를 활용하여 여러 회귀 모델을 학습하고, 성능 비교를 통해 최종 모델을 선택한 뒤 "
        "Streamlit 앱으로 배포 가능한 형태까지 구현하였다. 제출 산출물은 모델 학습 노트북, 한글 보고서, Streamlit 앱으로 구성된다.",
    )
    add_bullet(doc, "추가 데이터가 확보되면 제조사, 차량명, 국가(origin) 등 범주형 변수를 포함할 수 있다.")
    add_bullet(doc, "하이퍼파라미터 튜닝을 추가하면 모델 성능을 더 개선할 수 있다.")
    add_bullet(doc, "클라우드 DB를 연결하면 Streamlit Cloud에서도 예측 기록 저장 기능을 사용할 수 있다.")

    add_heading(doc, "9. 참고 자료", 1)
    add_bullet(doc, "UCI Machine Learning Repository, Auto MPG Dataset")
    add_bullet(doc, "Seaborn Data Repository, mpg.csv")
    add_bullet(doc, "scikit-learn Documentation")
    add_bullet(doc, "Streamlit Documentation")

    try:
        doc.save(REPORT_PATH)
        return REPORT_PATH
    except PermissionError:
        doc.save(FALLBACK_REPORT_PATH)
        return FALLBACK_REPORT_PATH


def main():
    df = load_data()
    trained_models, predictions, results, best_model_name, x_train, x_test, y_test = train_models(df)
    importance = get_feature_importance(trained_models[best_model_name])
    charts = create_charts(df, predictions, results, best_model_name, y_test, importance)
    saved_path = create_report(df, results, best_model_name, x_train, x_test, charts, importance)
    print(saved_path)


if __name__ == "__main__":
    main()
