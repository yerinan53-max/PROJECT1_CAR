# -*- coding: utf-8 -*-
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sklearn.linear_model import LinearRegression
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
from sklearn.model_selection import train_test_split


BASE = Path(__file__).resolve().parent
ASSETS = BASE / "report_assets"
OUTDIR = Path(r"C:/Users/human-21/Documents/Codex/2026-06-18/new-chat/outputs")
REPORT = OUTDIR / "car_mpg_ai_report_korean_fixed.docx"

FEATURES = [
    "cylinders",
    "displacement",
    "horsepower",
    "weight",
    "acceleration",
    "model_year",
]

LABELS = {
    "cylinders": "실린더 수",
    "displacement": "배기량",
    "horsepower": "마력",
    "weight": "무게",
    "acceleration": "가속력",
    "model_year": "연식",
}


def style_run(run, size=10.5, bold=False):
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    style_run(run, 9.5, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            set_cell_text(cell, value, row_index == 0 or col_index == 0)
            if row_index == 0 or col_index == 0:
                set_cell_shading(cell, "D9EAF7")
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        style_run(run, 16 if level == 1 else 13, True)
        run.font.color.rgb = RGBColor(22, 32, 51)
    return paragraph


def add_paragraph(doc, text=""):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    style_run(run, 10.5)
    paragraph.paragraph_format.line_spacing = 1.35
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    style_run(run, 10)
    return paragraph


def prepare_charts(df, model, y_test, y_pred, coef):
    ASSETS.mkdir(exist_ok=True)
    plt.rcParams["font.family"] = ["Malgun Gothic"]
    plt.rcParams["axes.unicode_minus"] = False

    fig, ax = plt.subplots(figsize=(7, 4))
    ax.hist(df["mpg"], bins=18, color="#1f7a8c", edgecolor="white")
    ax.set_title("자동차 연비(MPG) 분포")
    ax.set_xlabel("MPG")
    ax.set_ylabel("빈도")
    fig.tight_layout()
    mpg_hist = ASSETS / "mpg_distribution_fixed.png"
    fig.savefig(mpg_hist, dpi=160)
    plt.close(fig)

    corr = df[["mpg", *FEATURES]].corr()
    fig, ax = plt.subplots(figsize=(7, 5))
    image = ax.imshow(corr, cmap="RdBu_r", vmin=-1, vmax=1)
    ax.set_xticks(range(len(corr.columns)))
    ax.set_yticks(range(len(corr.columns)))
    ax.set_xticklabels(["mpg", *[LABELS[item] for item in FEATURES]], rotation=45, ha="right")
    ax.set_yticklabels(["mpg", *[LABELS[item] for item in FEATURES]])
    for row in range(len(corr.columns)):
        for col in range(len(corr.columns)):
            ax.text(col, row, f"{corr.iloc[row, col]:.2f}", ha="center", va="center", fontsize=8)
    fig.colorbar(image, ax=ax, shrink=0.8)
    ax.set_title("변수 간 상관관계 Heatmap")
    fig.tight_layout()
    heatmap = ASSETS / "correlation_heatmap_fixed.png"
    fig.savefig(heatmap, dpi=160)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(6.5, 4.5))
    ax.scatter(y_test, y_pred, alpha=0.75, color="#1769aa")
    min_value = min(y_test.min(), y_pred.min())
    max_value = max(y_test.max(), y_pred.max())
    ax.plot([min_value, max_value], [min_value, max_value], "--", color="#b3261e")
    ax.set_title("실제 연비와 예측 연비 비교")
    ax.set_xlabel("실제 MPG")
    ax.set_ylabel("예측 MPG")
    fig.tight_layout()
    prediction_plot = ASSETS / "actual_vs_predicted_fixed.png"
    fig.savefig(prediction_plot, dpi=160)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(7, 4))
    ax.barh(coef["label"], coef["absolute"], color="#1f7a8c")
    ax.invert_yaxis()
    ax.set_title("회귀 계수 기반 특성 영향도")
    ax.set_xlabel("계수 절댓값")
    fig.tight_layout()
    coef_plot = ASSETS / "feature_importance_fixed.png"
    fig.savefig(coef_plot, dpi=160)
    plt.close(fig)

    return mpg_hist, heatmap, prediction_plot, coef_plot


def main():
    OUTDIR.mkdir(parents=True, exist_ok=True)

    df = pd.read_csv(BASE / "data" / "mpg.csv")
    df = df[["mpg", *FEATURES]].dropna()

    X = df[FEATURES]
    y = df["mpg"]
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.2, random_state=42
    )
    model = LinearRegression()
    model.fit(X_train, y_train)
    y_pred = model.predict(X_test)

    metrics = {
        "R2": r2_score(y_test, y_pred),
        "MAE": mean_absolute_error(y_test, y_pred),
        "RMSE": np.sqrt(mean_squared_error(y_test, y_pred)),
    }
    coef = (
        pd.DataFrame(
            {
                "feature": FEATURES,
                "label": [LABELS[item] for item in FEATURES],
                "coefficient": model.coef_,
                "absolute": np.abs(model.coef_),
            }
        )
        .sort_values("absolute", ascending=False)
    )

    mpg_hist, heatmap, prediction_plot, coef_plot = prepare_charts(
        df, model, y_test, y_pred, coef
    )

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("2026년도 ML 모델 개발 과제\nAI개발 수행내역서")
    style_run(run, 22, True)

    doc.add_paragraph("")
    add_table(
        doc,
        [
            ["과제명", "자동차 제원 데이터를 활용한 연비 예측 모델 개발 및 시각화"],
            ["담당자", ""],
            ["작성일", "2026년 6월 22일"],
        ],
    )
    doc.add_paragraph("")
    add_paragraph(
        doc,
        "본 문서는 자동차의 주요 제원 데이터를 활용하여 연비(MPG)를 예측하는 "
        "선형회귀 기반 머신러닝 모델을 개발하고, Flask 웹 화면 및 MariaDB "
        "연동 구조를 통해 예측 결과를 시연할 수 있도록 정리한 수행내역서이다.",
    )
    doc.add_page_break()

    add_heading(doc, "AI개발 수행내용", 1)
    add_heading(doc, "1. 사업과제", 2)
    add_paragraph(doc, "자동차 제원 데이터를 활용한 AI 기반 연비 예측 모델 개발 및 웹 시각화")

    add_heading(doc, "2. 개요 및 현황", 2)
    add_heading(doc, "2.1 추진배경 및 목적", 3)
    for item in [
        "자동차 연비는 차량 유지비, 에너지 소비량, 환경 부담과 직접적으로 연결되는 중요한 지표이다.",
        "차량의 무게, 마력, 배기량, 실린더 수 등은 연비와 밀접한 관련이 있으므로 데이터를 기반으로 연비를 예측할 수 있다.",
        "본 프로젝트는 Auto MPG 데이터셋을 활용하여 선형회귀 모델을 학습하고, 사용자가 차량 제원을 입력하면 예상 연비를 웹 화면에서 확인할 수 있도록 구현하는 것을 목적으로 한다.",
        "추가로 MariaDB를 연결하여 예측 입력값과 결과를 저장함으로써 단순 예측 모델을 데이터 관리가 가능한 웹 시스템으로 확장하였다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.2 과제 범위", 3)
    add_table(
        doc,
        [
            ["과제구분", "내용"],
            ["AI 모델", "자동차 제원 기반 연비(MPG) 예측 선형회귀 모델 구현"],
            ["데이터", "Auto MPG 공개 데이터셋 수집, 결측치 제거, 독립변수/종속변수 분리"],
            ["모델 평가", "R², MAE, RMSE 지표를 활용한 성능 평가"],
            ["웹 프로토타입", "Flask 기반 입력 화면 및 예측 결과 출력 화면 구현"],
            ["DB 연동", "MariaDB에 예측 입력값과 예측 결과 저장 구조 구현"],
        ],
    )

    add_heading(doc, "2.3 과제 추진 방법", 3)
    add_paragraph(
        doc,
        "본 과제는 데이터 수집, 데이터 전처리, 모델 학습, 성능 평가, 웹 프로토타입 구현, DB 연동 순서로 진행하였다.",
    )
    add_table(
        doc,
        [
            ["단계", "수행 내용"],
            ["DATA IMPORTING", "mpg.csv 데이터 로드 및 주요 변수 확인"],
            ["USER INSIGHT", "연비와 차량 제원 간 상관관계 및 회귀 계수 분석"],
            ["AUTO PREDICTION", "Flask 화면에서 입력한 차량 제원을 기반으로 예상 MPG 예측"],
        ],
    )
    doc.add_page_break()

    add_heading(doc, "연구개발 주요 결과물", 1)
    add_heading(doc, "1. 데이터 수집", 2)
    add_paragraph(
        doc,
        "본 프로젝트에서는 Seaborn 예제 데이터로 제공되는 Auto MPG 데이터셋을 사용하였다. "
        "원 데이터셋은 UCI Machine Learning Repository의 Auto MPG Dataset이며, 자동차의 연비와 차량 제원 정보를 포함한다.",
    )
    add_table(
        doc,
        [
            ["항목", "내용"],
            ["데이터명", "Auto MPG Dataset"],
            ["출처", "UCI Machine Learning Repository / Seaborn mpg.csv"],
            ["데이터 개수", f"{len(df)}개"],
            ["종속변수(Y)", "mpg: 자동차 연비"],
            ["독립변수(X)", "cylinders, displacement, horsepower, weight, acceleration, model_year"],
        ],
    )

    add_heading(doc, "1.1 변수 구성", 3)
    add_table(
        doc,
        [
            ["변수명", "설명", "구분"],
            ["mpg", "자동차 연비(Miles Per Gallon)", "종속변수"],
            ["cylinders", "엔진 실린더 수", "독립변수"],
            ["displacement", "배기량", "독립변수"],
            ["horsepower", "마력", "독립변수"],
            ["weight", "차량 무게", "독립변수"],
            ["acceleration", "가속력", "독립변수"],
            ["model_year", "차량 연식", "독립변수"],
        ],
    )

    add_heading(doc, "2. 데이터 분석 및 전처리", 2)
    add_heading(doc, "2.1 기초 통계", 3)
    rows = [["변수", "평균", "최소값", "최대값"]]
    for column in ["mpg", *FEATURES]:
        rows.append(
            [
                LABELS.get(column, column),
                f"{df[column].mean():.2f}",
                f"{df[column].min():.2f}",
                f"{df[column].max():.2f}",
            ]
        )
    add_table(doc, rows)
    add_paragraph(
        doc,
        "데이터셋에서 결측치가 있는 행은 제거하였다. 특히 horsepower 변수는 일부 결측값이 존재할 수 있으므로 "
        "모델 학습 전에 dropna()를 적용하여 학습 가능한 수치형 데이터만 사용하였다.",
    )
    doc.add_picture(str(mpg_hist), width=Inches(5.8))
    add_paragraph(doc, "연비 데이터는 약 10~45 MPG 범위에 분포하며, 대부분의 차량은 중간 연비 구간에 집중되어 있다.")
    doc.add_picture(str(heatmap), width=Inches(5.8))
    add_paragraph(
        doc,
        "상관관계 분석 결과 차량 무게, 배기량, 마력, 실린더 수는 연비와 음의 상관관계를 보였다. "
        "즉 차량이 무겁고 엔진 규모가 클수록 연비가 낮아지는 경향이 나타났다.",
    )
    doc.add_page_break()

    add_heading(doc, "3. 모델 학습 및 평가", 2)
    add_heading(doc, "3.1 모델 정의", 3)
    add_paragraph(
        doc,
        "본 프로젝트에서는 scikit-learn의 LinearRegression 모델을 사용하였다. 선형회귀는 독립변수와 "
        "종속변수 사이의 선형 관계를 학습하여 연속적인 숫자값을 예측하는 회귀 모델이다.",
    )
    add_table(
        doc,
        [
            ["항목", "내용"],
            ["모델", "LinearRegression"],
            ["학습 데이터 비율", "80%"],
            ["테스트 데이터 비율", "20%"],
            ["random_state", "42"],
            ["평가 지표", "R², MAE, RMSE"],
        ],
    )
    add_heading(doc, "3.2 모델 성능 평가", 3)
    add_table(
        doc,
        [
            ["평가지표", "값", "설명"],
            ["R²", f"{metrics['R2']:.3f}", "1에 가까울수록 모델 설명력이 높음"],
            ["MAE", f"{metrics['MAE']:.3f}", "실제값과 예측값 차이의 평균"],
            ["RMSE", f"{metrics['RMSE']:.3f}", "큰 오차에 더 민감한 평균 제곱근 오차"],
        ],
    )
    doc.add_picture(str(prediction_plot), width=Inches(5.6))
    add_paragraph(
        doc,
        "실제 연비와 예측 연비의 산점도를 보면 대체로 대각선 근처에 분포하여 모델이 전반적인 연비 추세를 "
        "학습했음을 확인할 수 있다. 다만 모든 차량의 연비를 완벽히 맞추는 것은 아니며, 단순 선형회귀 모델의 한계가 존재한다.",
    )

    add_heading(doc, "3.3 특성 영향 분석", 3)
    rows = [["특성", "회귀 계수", "해석"]]
    for _, row in coef.iterrows():
        rows.append(
            [
                row["label"],
                f"{row['coefficient']:.4f}",
                "연비 증가 방향" if row["coefficient"] > 0 else "연비 감소 방향",
            ]
        )
    add_table(doc, rows)
    doc.add_picture(str(coef_plot), width=Inches(5.8))
    add_paragraph(
        doc,
        "회귀 계수의 부호와 크기를 통해 각 변수가 연비 예측에 미치는 방향을 확인할 수 있다. "
        "특히 차량 무게와 배기량은 연비와 밀접한 관계를 보이며, 차량이 무거울수록 연비가 낮아지는 경향이 나타났다.",
    )
    doc.add_page_break()

    add_heading(doc, "4. 프로토타입 구현", 2)
    add_heading(doc, "4.1 Flask 웹 화면", 3)
    add_paragraph(
        doc,
        "Flask를 활용하여 사용자가 자동차 제원 값을 입력하고 예측 버튼을 누르면 예상 연비를 확인할 수 있는 웹 화면을 구현하였다. "
        "입력 항목은 실린더 수, 배기량, 마력, 무게, 가속력, 연식으로 구성하였다.",
    )
    screenshot = ASSETS / "flask_screen.png"
    if screenshot.exists():
        doc.add_picture(str(screenshot), width=Inches(6.0))
    add_paragraph(
        doc,
        "웹 화면은 모델 성능 지표와 특성 영향 분석 결과를 함께 보여주도록 구성하여, 단순 예측 결과뿐 아니라 모델 해석 정보도 확인할 수 있게 하였다.",
    )

    add_heading(doc, "4.2 MariaDB 연동", 3)
    add_paragraph(
        doc,
        "MariaDB는 예측 요청 기록을 저장하기 위해 사용하였다. 사용자가 웹 화면에서 예측 버튼을 누르면 "
        "입력값과 예측 MPG가 predictions 테이블에 저장되도록 설계하였다.",
    )
    add_table(
        doc,
        [
            ["구분", "내용"],
            ["DBMS", "MariaDB"],
            ["Database", "project1_car"],
            ["Table", "predictions"],
            ["저장 데이터", "입력 차량 제원, 예측 MPG, 생성 시간"],
            ["연결 방식", "PyMySQL 라이브러리를 이용한 Python-MariaDB 연결"],
        ],
    )
    add_paragraph(doc, "보안을 위해 MariaDB 비밀번호는 코드에 직접 작성하지 않고 환경변수 DB_PASSWORD를 통해 입력하도록 구성하였다.")

    add_heading(doc, "5. 결론", 2)
    for item in [
        "본 프로젝트에서는 자동차 제원 데이터를 활용하여 연비를 예측하는 선형회귀 모델을 구현하였다.",
        f"모델 성능은 R²={metrics['R2']:.3f}, MAE={metrics['MAE']:.3f}, RMSE={metrics['RMSE']:.3f}로 나타났으며, 기본적인 연비 예측에는 활용 가능한 수준의 결과를 보였다.",
        "상관관계와 회귀 계수 분석을 통해 차량 무게, 배기량, 마력 등이 연비와 관련이 있음을 확인하였다.",
        "Flask 웹 화면과 MariaDB 저장 구조를 추가하여 사용자가 직접 값을 입력하고 예측 결과를 확인할 수 있는 시연 가능한 프로토타입을 구축하였다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. 향후 발전 방향", 2)
    for item in [
        "선형회귀 외에 RandomForestRegressor, XGBoost 등 다양한 회귀 모델을 비교하여 예측 성능을 높일 수 있다.",
        "제조사, 차량명, 국가(origin) 등 범주형 변수를 추가하여 더 풍부한 예측 모델을 만들 수 있다.",
        "MariaDB에 저장된 예측 기록을 활용하여 사용자 입력 패턴이나 예측 이력을 시각화하는 기능을 추가할 수 있다.",
        "실제 최신 자동차 제원 데이터를 추가 수집하면 과제용 예제 수준을 넘어 현실적인 연비 예측 서비스로 확장할 수 있다.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "참고자료", 1)
    for item in [
        "UCI Machine Learning Repository, Auto MPG Dataset: https://archive.ics.uci.edu/dataset/9/auto+mpg",
        "Seaborn Data Repository, mpg.csv: https://raw.githubusercontent.com/mwaskom/seaborn-data/master/mpg.csv",
        "Flask Documentation: https://flask.palletsprojects.com/",
        "scikit-learn Documentation: https://scikit-learn.org/",
        "MariaDB Documentation: https://mariadb.org/documentation/",
    ]:
        add_bullet(doc, item)

    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    main()
